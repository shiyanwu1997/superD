from __future__ import annotations

import math
from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib import font_manager
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "SuperD软件著作权登记说明书_V1.0.docx"
IMG_DIR = ROOT / "_softcopyright_assets"
IMG_DIR.mkdir(parents=True, exist_ok=True)

SOFTWARE_NAME = "SuperD Supervisor进程管理平台软件"
VERSION = "V1.0"
RIGHTSHOLDER = "北京胤康科技有限公司"
GEN_DATE = "2026年8月28日"

BLUE = "1F4D78"
MID_BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
INK = "172B3A"
MUTED = "5F6B76"
WHITE = "FFFFFF"
ACCENT = "2F6B5F"
GOLD = "8A6A18"
RED = "9B1C1C"

# Preset: compact_reference_guide.
# Named overrides: A4_softcopyright_submission (A4 portrait, 9360 DXA content),
# CJK_typography (SimSun body, Microsoft YaHei headings), and
# Chinese_registration_furniture (software/version header and continuous PAGE field).
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def chinese_font_path() -> str | None:
    candidates = [
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return candidate
    return None


FONT_PROP = font_manager.FontProperties(fname=chinese_font_path()) if chinese_font_path() else None


def set_run_font(run, name="宋体", size=11, bold=None, color=INK, italic=None):
    run.font.name = name
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths, indent=TABLE_INDENT_DXA):
    if sum(widths) != CONTENT_DXA:
        raise ValueError(f"Table widths must sum to {CONTENT_DXA}: {widths}")
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    old_grid = tbl.tblGrid
    new_grid = OxmlElement("w:tblGrid")
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        new_grid.append(grid_col)
    tbl.replace(old_grid, new_grid)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            set_cell_margins(cell, **CELL_MARGINS)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])
    set_run_font(run, "宋体", 9, color=MUTED)


def set_page_number_start(section, start=1):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))


def add_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_id = 30
    num_id = 30
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.extend([tabs, ind, spacing])
    lvl.extend([start, num_fmt, lvl_text, suff, ppr])
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)

    bullet_abstract_id = 31
    bullet_num_id = 31
    abstract_b = OxmlElement("w:abstractNum")
    abstract_b.set(qn("w:abstractNumId"), str(bullet_abstract_id))
    abstract_b.append(parse_xml(f'<w:multiLevelType {nsdecls("w")} w:val="singleLevel"/>'))
    lvl_b = OxmlElement("w:lvl")
    lvl_b.set(qn("w:ilvl"), "0")
    fmt_b = OxmlElement("w:numFmt")
    fmt_b.set(qn("w:val"), "bullet")
    text_b = OxmlElement("w:lvlText")
    text_b.set(qn("w:val"), "●")
    ppr_b = OxmlElement("w:pPr")
    tabs_b = OxmlElement("w:tabs")
    tab_b = OxmlElement("w:tab")
    tab_b.set(qn("w:val"), "num")
    tab_b.set(qn("w:pos"), "540")
    tabs_b.append(tab_b)
    ind_b = OxmlElement("w:ind")
    ind_b.set(qn("w:left"), "540")
    ind_b.set(qn("w:hanging"), "270")
    spacing_b = OxmlElement("w:spacing")
    spacing_b.set(qn("w:after"), "80")
    spacing_b.set(qn("w:line"), "300")
    spacing_b.set(qn("w:lineRule"), "auto")
    ppr_b.extend([tabs_b, ind_b, spacing_b])
    lvl_b.extend([fmt_b, text_b, ppr_b])
    abstract_b.append(lvl_b)
    numbering.append(abstract_b)
    num_b = OxmlElement("w:num")
    num_b.set(qn("w:numId"), str(bullet_num_id))
    abstract_num_id_b = OxmlElement("w:abstractNumId")
    abstract_num_id_b.set(qn("w:val"), str(bullet_abstract_id))
    num_b.append(abstract_num_id_b)
    numbering.append(num_b)
    return num_id, bullet_num_id


def apply_num(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    ppr.append(num_pr)


def style_document(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in [
        ("Heading 1", 16, MID_BLUE, 18, 10),
        ("Heading 2", 13, MID_BLUE, 14, 7),
        ("Heading 3", 12, BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "PingFang SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    title = styles["Title"]
    title.font.name = "PingFang SC"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(BLUE)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_ppr = title._element.get_or_add_pPr()
    title_border = title_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_ppr.remove(title_border)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "PingFang SC"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    subtitle.font.size = Pt(15)
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def configure_section(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Inches(1273 / 1440)
    section.right_margin = Inches(1273 / 1440)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)


def setup_running_furniture(section):
    section.header.is_linked_to_previous = False
    header = section.header
    p = header.paragraphs[0]
    p.text = ""
    p.paragraph_format.space_after = Pt(2)
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [7100, 2260], indent=0)
    table.style = "Table Grid"
    for cell in table.rows[0].cells:
        set_cell_shading(cell, WHITE)
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        for edge in ("top", "left", "right", "insideH", "insideV"):
            edge_el = OxmlElement(f"w:{edge}")
            edge_el.set(qn("w:val"), "nil")
            borders.append(edge_el)
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:color"), "B9C7D3")
        borders.append(bottom)
    left = table.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(left.add_run(f"{SOFTWARE_NAME} {VERSION}"), "宋体", 9, bold=True, color=MUTED)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(right.add_run("第 "), "宋体", 9, color=MUTED)
    add_page_field(right)
    set_run_font(right.add_run(" 页"), "宋体", 9, color=MUTED)

    section.footer.is_linked_to_previous = False
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(3)
    set_run_font(fp.add_run(f"{RIGHTSHOLDER} | 软件著作权登记技术文档"), "宋体", 8.5, color=MUTED)


def add_body(doc, text, bold_lead=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.first_line_indent = Cm(0.74)
    if bold_lead and text.startswith(bold_lead):
        set_run_font(p.add_run(bold_lead), "宋体", 11, bold=True, color=INK)
        set_run_font(p.add_run(text[len(bold_lead):]), "宋体", 11, color=INK)
    else:
        set_run_font(p.add_run(text), "宋体", 11, color=INK)
    return p


def add_note(doc, label, text, fill=LIGHT_BLUE, color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(f"{label}："), "PingFang SC", 10.5, bold=True, color=color)
    set_run_font(p.add_run(text), "宋体", 10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_list(doc, items, numbered=False):
    global NEXT_NUM_ID
    if numbered:
        NEXT_NUM_ID += 1
        num_id = NEXT_NUM_ID
        numbering = doc.part.numbering_part.element
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_num_id = OxmlElement("w:abstractNumId")
        abstract_num_id.set(qn("w:val"), "30")
        num.append(abstract_num_id)
        lvl_override = OxmlElement("w:lvlOverride")
        lvl_override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        lvl_override.append(start_override)
        num.append(lvl_override)
        numbering.append(num)
    else:
        num_id = BULLET_NUM_ID
    for item in items:
        p = doc.add_paragraph()
        apply_num(p, num_id)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        set_run_font(p.add_run(item), "宋体", 11, color=INK)


def add_table(doc, headers, rows, widths, font_size=9.5, header_fill=LIGHT_BLUE, aligns=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, value in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        set_run_font(p.add_run(str(value)), "PingFang SC", font_size, bold=True, color=BLUE)
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            if row_idx % 2 == 1:
                set_cell_shading(cell, "FAFBFC")
            p = cell.paragraphs[0]
            p.alignment = aligns[i] if aligns else (WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            set_run_font(p.add_run(str(value)), "宋体", font_size, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_figure(doc, image_path, caption, width=6.15):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(image_path), width=Inches(width))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after = Pt(8)
    cp.paragraph_format.keep_with_next = False
    set_run_font(cp.add_run(caption), "宋体", 9.5, color=MUTED)


def draw_box(ax, xy, width, height, title, subtitle="", fill="#E8EEF5", edge="#2E74B5"):
    x, y = xy
    patch = FancyBboxPatch((x, y), width, height, boxstyle="round,pad=0.012,rounding_size=0.02",
                           linewidth=1.2, edgecolor=edge, facecolor=fill)
    ax.add_patch(patch)
    ax.text(x + width / 2, y + height * 0.63, title, ha="center", va="center",
            fontsize=11, fontweight="bold", color="#1F4D78", fontproperties=FONT_PROP)
    if subtitle:
        ax.text(x + width / 2, y + height * 0.28, subtitle, ha="center", va="center",
                fontsize=8.2, color="#44515D", fontproperties=FONT_PROP, wrap=True)


def arrow(ax, start, end, text=""):
    ax.add_patch(FancyArrowPatch(start, end, arrowstyle="-|>", mutation_scale=12,
                                 linewidth=1.1, color="#617889"))
    if text:
        ax.text((start[0] + end[0]) / 2, (start[1] + end[1]) / 2 + 0.025, text,
                ha="center", va="center", fontsize=7.7, color="#526473", fontproperties=FONT_PROP)


def save_diagrams():
    plt.rcParams["axes.unicode_minus"] = False
    # Architecture.
    fig, ax = plt.subplots(figsize=(10.5, 5.7), dpi=170)
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    draw_box(ax, (0.03, 0.68), 0.20, 0.20, "浏览器客户端", "React + Ant Design\nTanStack Query / xterm.js", "#F2F4F7")
    draw_box(ax, (0.30, 0.68), 0.22, 0.20, "Web服务层", "Express 5 / JWT / API Token\nHelmet / Rate Limit", "#E8EEF5")
    draw_box(ax, (0.59, 0.68), 0.18, 0.20, "实时通信", "Socket.io\n日志增量推送", "#EAF3EE", "#2F6B5F")
    draw_box(ax, (0.30, 0.28), 0.22, 0.20, "数据访问层", "SQLite（默认）/ MySQL\n用户、权限、项目、审计", "#FFF8E6", "#8A6A18")
    draw_box(ax, (0.68, 0.24), 0.27, 0.26, "Supervisor节点集群", "XML-RPC :9001\n进程状态、启停、日志、重载\n支持多服务器接入", "#F4EDED", "#9B1C1C")
    arrow(ax, (0.23, 0.78), (0.30, 0.78), "HTTPS / JSON")
    arrow(ax, (0.52, 0.78), (0.59, 0.78), "WebSocket")
    arrow(ax, (0.41, 0.68), (0.41, 0.48), "数据读写")
    arrow(ax, (0.52, 0.72), (0.70, 0.50), "XML-RPC")
    arrow(ax, (0.77, 0.68), (0.82, 0.50), "日志轮询")
    ax.text(0.5, 0.95, "SuperD系统总体架构", ha="center", fontsize=15, fontweight="bold",
            color="#172B3A", fontproperties=FONT_PROP)
    fig.tight_layout(); fig.savefig(IMG_DIR / "architecture.png", bbox_inches="tight", facecolor="white"); plt.close(fig)

    # RBAC.
    fig, ax = plt.subplots(figsize=(10.5, 5.6), dpi=170)
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    draw_box(ax, (0.37, 0.72), 0.26, 0.16, "超级管理员 admin", "全局用户、机器、分组、令牌管理", "#E8EEF5")
    draw_box(ax, (0.11, 0.42), 0.28, 0.16, "普通管理员 subadmin", "管理自己的组员并分配权限", "#EAF3EE", "#2F6B5F")
    draw_box(ax, (0.61, 0.42), 0.28, 0.16, "普通用户 user", "按授权查看并操作进程", "#F2F4F7")
    draw_box(ax, (0.18, 0.12), 0.27, 0.14, "机器级授权", "可操作该机器全部程序", "#FFF8E6", "#8A6A18")
    draw_box(ax, (0.55, 0.12), 0.27, 0.14, "程序级授权", "仅操作指定程序，优先判定", "#FFF8E6", "#8A6A18")
    arrow(ax, (0.44, 0.72), (0.28, 0.58), "创建/管理")
    arrow(ax, (0.56, 0.72), (0.72, 0.58), "直接管理")
    arrow(ax, (0.29, 0.42), (0.30, 0.26), "分配")
    arrow(ax, (0.72, 0.42), (0.68, 0.26), "使用")
    ax.text(0.5, 0.95, "角色与资源权限模型", ha="center", fontsize=15, fontweight="bold",
            color="#172B3A", fontproperties=FONT_PROP)
    fig.tight_layout(); fig.savefig(IMG_DIR / "rbac.png", bbox_inches="tight", facecolor="white"); plt.close(fig)

    # Process operation sequence.
    fig, ax = plt.subplots(figsize=(10.5, 5.7), dpi=170)
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    xs = [0.10, 0.35, 0.60, 0.85]
    titles = ["用户界面", "权限中间件", "服务与缓存", "Supervisor"]
    for x, title in zip(xs, titles):
        ax.text(x, 0.92, title, ha="center", fontsize=10.5, fontweight="bold", color="#1F4D78", fontproperties=FONT_PROP)
        ax.plot([x, x], [0.12, 0.86], linestyle="--", linewidth=0.8, color="#AAB8C2")
    events = [
        (0.82, 0, 1, "提交 start / stop / restart"),
        (0.68, 1, 2, "验证角色、机器/程序权限"),
        (0.54, 2, 3, "XML-RPC 执行控制指令"),
        (0.40, 3, 2, "返回执行结果"),
        (0.28, 2, 2, "清除5秒进程缓存"),
        (0.17, 2, 0, "返回成功并触发状态刷新"),
    ]
    for y, src, dst, label in events:
        if src == dst:
            ax.add_patch(FancyArrowPatch((xs[src], y), (xs[src] + 0.10, y - 0.045),
                                         connectionstyle="arc3,rad=-0.55", arrowstyle="-|>",
                                         mutation_scale=10, linewidth=1.0, color="#617889"))
            ax.text(xs[src] + 0.13, y - 0.02, label, fontsize=8, color="#526473", fontproperties=FONT_PROP)
        else:
            direction = 1 if dst > src else -1
            start_x = xs[src] + 0.015 * direction
            end_x = xs[dst] - 0.015 * direction
            arrow(ax, (start_x, y), (end_x, y), label)
    ax.text(0.5, 0.98, "进程控制请求时序", ha="center", fontsize=15, fontweight="bold",
            color="#172B3A", fontproperties=FONT_PROP)
    fig.tight_layout(); fig.savefig(IMG_DIR / "sequence.png", bbox_inches="tight", facecolor="white"); plt.close(fig)

    # Deployment.
    fig, ax = plt.subplots(figsize=(10.5, 5.5), dpi=170)
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    draw_box(ax, (0.03, 0.62), 0.20, 0.20, "运维用户", "Chrome / Edge\n企业内网访问", "#F2F4F7")
    draw_box(ax, (0.31, 0.55), 0.27, 0.32, "SuperD应用服务器", "PM2托管 Node.js\n单端口6002提供Web与API\n环境变量注入密钥", "#E8EEF5")
    draw_box(ax, (0.34, 0.15), 0.21, 0.18, "持久化存储", "SQLite文件\n或独立MySQL", "#FFF8E6", "#8A6A18")
    draw_box(ax, (0.68, 0.60), 0.27, 0.23, "被管服务器 A...N", "Supervisor inet_http_server\nXML-RPC :9001", "#EAF3EE", "#2F6B5F")
    draw_box(ax, (0.68, 0.20), 0.27, 0.19, "日志与审计", "应用访问日志 / 错误日志\n服务API调用审计", "#F4EDED", "#9B1C1C")
    arrow(ax, (0.23, 0.72), (0.31, 0.72), "HTTPS")
    arrow(ax, (0.58, 0.72), (0.68, 0.72), "XML-RPC")
    arrow(ax, (0.44, 0.55), (0.44, 0.33), "读写")
    arrow(ax, (0.58, 0.60), (0.68, 0.38), "记录")
    ax.text(0.5, 0.95, "推荐生产部署拓扑", ha="center", fontsize=15, fontweight="bold",
            color="#172B3A", fontproperties=FONT_PROP)
    fig.tight_layout(); fig.savefig(IMG_DIR / "deployment.png", bbox_inches="tight", facecolor="white"); plt.close(fig)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    set_run_font(p.add_run(text), "PingFang SC", {1: 16, 2: 13, 3: 12}[level], bold=True,
                 color={1: MID_BLUE, 2: MID_BLUE, 3: BLUE}[level])
    return p


def build():
    save_diagrams()
    doc = Document()
    style_document(doc)
    global DECIMAL_NUM_ID, BULLET_NUM_ID, NEXT_NUM_ID
    DECIMAL_NUM_ID, BULLET_NUM_ID = add_numbering(doc)
    NEXT_NUM_ID = 40
    configure_section(doc.sections[0])

    # Cover: editorial_cover header pattern adapted for a formal software manual.
    cover = doc.sections[0]
    cover.header.is_linked_to_previous = False
    cover.footer.is_linked_to_previous = False
    cover.header.paragraphs[0].text = ""
    cover.footer.paragraphs[0].text = ""
    doc.add_paragraph().paragraph_format.space_after = Pt(58)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(kicker.add_run("计算机软件著作权登记技术文档"), "PingFang SC", 13, bold=True, color=GOLD)
    title = doc.add_paragraph(style="Title")
    title.add_run(SOFTWARE_NAME)
    title_ppr = title._p.get_or_add_pPr()
    title_border = title_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_ppr.remove(title_border)
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("软件说明书（设计说明书及用户操作手册）")
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version.paragraph_format.space_before = Pt(12)
    version.paragraph_format.space_after = Pt(86)
    set_run_font(version.add_run(VERSION), "PingFang SC", 18, bold=True, color=ACCENT)

    meta = doc.add_table(rows=5, cols=2)
    meta.style = "Table Grid"
    set_table_geometry(meta, [2700, 6660])
    cover_rows = [
        ("软件著作权人", RIGHTSHOLDER),
        ("软件版本号", VERSION),
        ("文档类别", "软件鉴别材料 / 技术说明文档"),
        ("编制日期", GEN_DATE),
        ("文档状态", "登记申请稿（主体信息待申请人核定）"),
    ]
    for idx, (label, value) in enumerate(cover_rows):
        set_cell_shading(meta.cell(idx, 0), LIGHT_BLUE)
        p0 = meta.cell(idx, 0).paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p0.add_run(label), "PingFang SC", 10.5, bold=True, color=BLUE)
        p1 = meta.cell(idx, 1).paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_run_font(p1.add_run(value), "宋体", 10.5, color=INK)
        for cell in meta.rows[idx].cells:
            set_cell_margins(cell, top=110, bottom=110, start=140, end=140)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(28)
    set_run_font(cp.add_run(RIGHTSHOLDER), "宋体", 12, bold=True, color=MUTED)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(body_section)
    set_page_number_start(body_section, 1)
    setup_running_furniture(body_section)

    add_heading(doc, "文档控制信息", 1)
    add_body(doc, "本说明书依据SuperD项目当前仓库中的前端、后端、数据库初始化脚本、接口说明和部署配置编制，用于说明软件的开发目的、总体结构、主要功能、数据设计、运行环境、操作方法和安全机制。文档内容只陈述能够从项目材料核实的技术事实；著作权人和开发完成日期等客观信息由申请人最终核定。")
    add_table(doc, ["版本", "日期", "编制依据", "说明"], [
        ("V1.0", GEN_DATE, "SuperD当前项目仓库", "首次形成软件著作权登记说明书"),
    ], [1200, 1800, 3000, 3360], font_size=9.5)
    add_note(doc, "交存说明", "本说明书是登记所需的“文档鉴别材料”，按A4纵向编排，页眉统一标注软件全称和版本号，正文页连续编号；不足60页提交全部。", fill="FFF8E6", color=GOLD)

    add_heading(doc, "内容目录", 1)
    toc_items = [
        "1 引言", "2 软件概述", "3 运行环境", "4 总体设计", "5 功能设计",
        "6 数据设计", "7 用户操作说明", "8 服务API使用说明", "9 安全设计",
        "10 异常处理与日志", "11 安装部署与维护", "12 测试与验收建议",
        "13 技术特点与独创性说明", "附录A 主要接口清单", "附录B 功能与源代码对应关系",
        "附录C 界面视图清单",
    ]
    add_list(doc, toc_items, numbered=False)

    add_heading(doc, "1 引言", 1)
    add_heading(doc, "1.1 编写目的", 2)
    add_body(doc, "本说明书面向软件著作权登记审查、企业知识产权管理、系统部署人员和最终用户，完整描述SuperD软件的设计与使用方式。通过本说明书，读者可以了解软件解决的问题、技术组成、功能边界、用户角色、主要业务流程、数据存储方式和部署条件，并能够按照操作步骤完成登录、进程监控、进程控制、实时日志查看、机器与分组管理、用户授权和服务API令牌管理。")
    add_heading(doc, "1.2 软件命名", 2)
    add_table(doc, ["项目", "拟登记内容"], [
        ("软件全称", SOFTWARE_NAME),
        ("软件简称", "SuperD"),
        ("版本号", VERSION),
        ("软件类别", "应用软件 / 运维管理软件"),
        ("开发团队", "由公司运维团队进行开发"),
    ], [2100, 7260])
    add_heading(doc, "1.3 术语和缩略语", 2)
    add_table(doc, ["术语", "说明"], [
        ("Supervisor", "用于类Unix系统的进程控制系统，通过XML-RPC接口提供进程查询和控制能力。"),
        ("XML-RPC", "基于HTTP和XML的远程过程调用协议，SuperD通过该协议连接被管理服务器。"),
        ("RBAC", "基于角色的访问控制，本软件设有超级管理员、普通管理员和普通用户。"),
        ("JWT", "网页用户登录后使用的短期身份令牌。"),
        ("API Token", "面向服务集成的可撤销令牌，令牌值以sd_开头并受scope范围约束。"),
        ("项目/机器", "系统中的受管Supervisor服务器记录，界面中以机器或项目节点展示。"),
        ("程序", "由Supervisor管理的单个进程实例。"),
    ], [1800, 7560])

    add_heading(doc, "2 软件概述", 1)
    add_note(doc, "开发团队", "本软件由北京胤康科技有限公司运维团队进行开发。")
    add_heading(doc, "2.1 开发背景", 2)
    add_body(doc, "在多服务器环境中，运维人员通常需要分别登录每台主机并运行supervisorctl命令查询或控制进程。随着服务器数量、业务程序数量和协作人员增加，传统命令行方式会产生入口分散、权限边界粗、状态难以汇总、日志查看不统一以及操作审计不足等问题。SuperD提供统一的Web管理入口，把多台服务器上的Supervisor进程组织为可搜索、可分组、可授权的资源，并在浏览器中完成监控、启停、重启和日志查看。")
    add_heading(doc, "2.2 建设目标", 2)
    add_list(doc, [
        "建立一个统一的多服务器Supervisor进程管理控制台。",
        "以机器、分组和程序为层级组织资源，降低定位目标进程的成本。",
        "提供程序级和机器级权限控制，满足不同运维人员的最小授权需求。",
        "通过实时日志推送和自动刷新，使用户及时掌握程序状态和运行输出。",
        "通过SQLite零配置模式和MySQL可选模式兼顾快速部署与生产扩展。",
        "为外部监控、告警或自动化平台提供带scope和审计能力的服务API。",
    ])
    add_heading(doc, "2.3 适用对象", 2)
    add_table(doc, ["用户类型", "典型职责", "主要使用范围"], [
        ("超级管理员", "平台初始化、全局配置和安全治理", "管理所有用户、分组、机器、权限和服务令牌"),
        ("普通管理员", "团队或业务组运维管理", "管理自己的普通用户并为其配置资源权限"),
        ("普通用户", "日常进程监控和故障处理", "查看并操作已授权机器或指定程序"),
        ("外部服务", "告警、发布、巡检或自动化集成", "在令牌scope与绑定用户权限交集内调用API"),
    ], [1800, 3000, 4560])
    add_heading(doc, "2.4 软件边界", 2)
    add_body(doc, "SuperD负责提供Web界面、身份认证、权限判定、资源配置、Supervisor XML-RPC调用、日志传输、状态缓存和服务API审计。它不替代Supervisor本身，不负责创建操作系统进程配置文件，也不直接执行服务器Shell命令。被管理服务器必须已经安装并运行Supervisor，并开放可由SuperD应用服务器访问的inet_http_server/XML-RPC接口。")

    add_heading(doc, "3 运行环境", 1)
    add_heading(doc, "3.1 开发环境", 2)
    add_table(doc, ["类别", "技术/工具", "用途"], [
        ("前端语言", "JavaScript、JSX、CSS", "实现浏览器端页面与交互"),
        ("前端框架", "React 18、Vite 7、Ant Design 5", "组件化界面、构建与UI基础能力"),
        ("前端组件", "React Router、TanStack Query、xterm.js、Socket.io-client", "路由、数据请求、终端日志、实时通信"),
        ("后端环境", "Node.js、Express 5、Socket.io 4", "REST API、静态页面托管和实时日志"),
        ("数据存储", "SQLite / MySQL", "持久化用户、角色、资源、权限、令牌和审计数据"),
        ("进程通信", "xmlrpc", "调用Supervisor XML-RPC接口"),
        ("部署工具", "PM2", "生产环境进程守护和日志管理"),
    ], [1500, 3600, 4260], font_size=9.2)
    add_heading(doc, "3.2 服务器端运行条件", 2)
    add_list(doc, [
        "操作系统：支持Node.js和PM2的Linux/Unix服务器，推荐64位Linux。",
        "运行时：与项目依赖兼容的Node.js版本，生产部署前应使用锁定文件完成依赖安装和构建验证。",
        "存储：默认使用本地SQLite文件；并发或集中化场景可配置MySQL。",
        "网络：应用服务器能够访问各被管服务器的Supervisor XML-RPC端口，默认端口为9001。",
        "安全：建议通过HTTPS反向代理对外提供服务，密钥仅通过环境变量或受控配置文件注入。",
    ])
    add_heading(doc, "3.3 客户端运行条件", 2)
    add_body(doc, "客户端为现代Web浏览器，不需要安装独立桌面程序。建议使用当前受支持版本的Chrome、Edge或其他兼容现代JavaScript、WebSocket和CSS布局的浏览器，桌面分辨率建议不低于1366×768。客户端应能访问SuperD服务地址，并允许同源Cookie、JWT请求和Socket.io连接。")
    add_heading(doc, "3.4 关键环境变量", 2)
    add_table(doc, ["变量", "必填", "默认值", "用途"], [
        ("SESSION_SECRET", "是", "无", "Express Session签名密钥"),
        ("JWT_SECRET", "是", "无", "网页用户JWT签名密钥"),
        ("ENCRYPTION_KEY", "是", "无", "Supervisor连接密码加密密钥"),
        ("STORAGE_TYPE", "否", "sqlite", "选择sqlite或mysql存储实现"),
        ("SQLITE_PATH", "否", "backend/data/supervisor.db", "SQLite数据库文件位置"),
        ("PORT", "否", "3000", "后端监听端口；生产配置示例为6002"),
        ("MYSQL_*", "MySQL模式", "按配置", "MySQL主机、端口、用户、密码、数据库等"),
    ], [2000, 1200, 2700, 3460], font_size=9.1)

    add_heading(doc, "4 总体设计", 1)
    add_heading(doc, "4.1 系统架构", 2)
    add_body(doc, "软件采用前后端分离的B/S结构。前端负责页面展示、资源树导航、状态统计、用户交互和日志终端；后端提供认证、权限校验、REST接口、WebSocket服务、数据库访问及Supervisor XML-RPC适配。生产构建后，后端可在同一端口托管前端静态文件，从而形成单一访问入口。")
    add_figure(doc, IMG_DIR / "architecture.png", "图4-1 SuperD系统总体架构")
    add_heading(doc, "4.2 模块划分", 2)
    add_table(doc, ["层级", "模块", "主要职责"], [
        ("表现层", "登录页、进程主面板、用户管理、项目管理、日志终端", "接收用户操作并展示资源、状态、统计、日志与反馈"),
        ("前端状态层", "AuthContext、TanStack Query、API封装", "管理登录态、请求生命周期、错误处理与自动刷新"),
        ("接口层", "auth、projects、groups、programs、users、apiTokens路由", "提供REST接口并执行参数校验和业务编排"),
        ("安全层", "JWT/API Token认证、角色与资源权限中间件", "验证身份、scope以及机器/程序访问权限"),
        ("服务层", "supervisorService、socketServer", "执行XML-RPC调用、缓存进程列表、增量读取并推送日志"),
        ("数据层", "db.sqlite、db.mysql", "为两类数据库提供统一的数据访问方法"),
    ], [1400, 3400, 4560], font_size=9.0)
    add_heading(doc, "4.3 数据流", 2)
    add_body(doc, "用户登录成功后，前端以JWT访问受保护接口。后端根据用户角色和项目/程序权限筛选可见资源；查询进程时，服务层通过XML-RPC从目标Supervisor读取状态，并在5秒内复用缓存结果。控制操作成功后立即清除相关缓存，下一次查询获得最新状态。日志页面通过Socket.io建立实时连接，服务端按偏移量读取stdout或stderr增量内容并推送到xterm.js终端。")
    add_heading(doc, "4.4 部署拓扑", 2)
    add_figure(doc, IMG_DIR / "deployment.png", "图4-2 推荐生产部署拓扑")

    add_heading(doc, "5 功能设计", 1)
    add_heading(doc, "5.1 身份认证", 2)
    add_body(doc, "登录接口接收用户名和密码，后端使用bcrypt校验密码摘要，成功后签发JWT并返回用户基本信息。前端AuthContext统一保存认证状态，受保护路由在未登录时跳转到登录页。登录请求设置按IP计数的速率限制，当前实现为每分钟最多5次，以降低暴力猜测风险。用户可以主动退出并清理会话信息。")
    add_heading(doc, "5.2 分组与机器管理", 2)
    add_body(doc, "超级管理员可维护项目分组和机器连接信息。机器记录包含名称、说明、所属分组以及Supervisor主机、端口、用户名和密码等配置。连接密码在写入数据库前使用AES-256相关加密逻辑处理。前端以“分组→机器→程序”的层级展示资源，并按机器名称前缀和尾部数字进行自然排序，支持关键字筛选和连接状态指示。")
    add_heading(doc, "5.3 进程监控", 2)
    add_body(doc, "用户选择机器后，系统从Supervisor获取该机器上的程序列表，并展示名称、状态、运行时间和操作入口。页面根据当前程序集合计算总数、运行中、已停止和异常数量。自动刷新周期为10秒，刷新时只更新状态和运行时间等变化字段，不显示全屏加载动画，避免打断用户操作。全局视图能够汇总用户有权访问的程序。")
    add_heading(doc, "5.4 进程控制", 2)
    add_body(doc, "对于单个程序，用户可执行启动、停止和重启；对于一台机器，可执行全部启动、全部停止、全部重启和Supervisor配置重载。每次请求先通过身份、scope、角色和资源权限校验，再调用对应XML-RPC方法。操作成功后服务层清除缓存，前端显示结果并延迟刷新状态。")
    add_figure(doc, IMG_DIR / "sequence.png", "图5-1 进程控制请求时序")
    add_heading(doc, "5.5 实时日志", 2)
    add_body(doc, "程序详情提供stdout和stderr日志查看能力。前端使用xterm.js模拟终端，支持暂停/继续、滚动和选中复制。Socket.io连接建立后，服务端先确定日志文件当前位置，再按偏移量增量读取新内容；连续无新日志时，轮询间隔自动从1秒降低到10秒，以减少服务器负载，出现新日志后恢复高频读取。单次读取长度和前端最大显示行数受配置限制。")
    add_heading(doc, "5.6 用户与权限管理", 2)
    add_body(doc, "系统设置超级管理员、普通管理员和普通用户三级角色。超级管理员可管理所有用户；普通管理员只能创建和管理自己的普通用户；普通用户不能进入用户管理功能。资源权限支持机器级和程序级两种粒度，程序级权限优先判定，机器级权限作为兜底。该设计既能让用户管理整台机器，也能只授权某个具体程序。")
    add_figure(doc, IMG_DIR / "rbac.png", "图5-2 角色与资源权限模型")
    add_heading(doc, "5.7 服务API令牌与审计", 2)
    add_body(doc, "系统支持面向长期服务调用的API令牌。令牌以sd_为前缀，创建时只返回一次明文，数据库仅保存SHA-256摘要。每个令牌绑定一个用户、名称、scope集合和有效期，可由管理员立即撤销。请求权限必须同时满足令牌scope和绑定用户的资源权限。服务令牌调用完成后，系统记录令牌、用户、方法、路径、状态码、耗时和IP等审计信息。")
    add_heading(doc, "5.8 性能与一致性", 2)
    add_list(doc, [
        "进程列表使用5秒缓存，降低频繁刷新对Supervisor的XML-RPC压力。",
        "控制操作完成后清除对应缓存，避免状态查询长期返回旧值。",
        "页面10秒静默刷新，只合并变化字段，减少布局抖动。",
        "机器连接状态检查采用有限重试，失败后展示离线或错误信息。",
        "日志连续为空时自动降频，有新日志后恢复高频读取。",
    ])

    add_heading(doc, "6 数据设计", 1)
    add_heading(doc, "6.1 存储模式", 2)
    add_body(doc, "数据访问层根据STORAGE_TYPE选择SQLite或MySQL实现。SQLite模式无需独立数据库服务，适合单机或快速部署；MySQL模式适合集中存储和生产环境。两种实现提供一致的核心数据访问方法，使上层路由和服务不依赖具体数据库类型。")
    add_heading(doc, "6.2 核心数据实体", 2)
    add_table(doc, ["实体", "关键字段", "用途"], [
        ("roles", "id、name、description", "保存admin、subadmin、user角色定义"),
        ("users", "username、password、roleId、createdBy", "保存用户、密码摘要、角色和上级管理员"),
        ("project_groups", "name、description", "对受管机器进行业务分组"),
        ("projects", "name、description、groupId、supervisorConfig", "保存Supervisor服务器连接配置"),
        ("user_project_permissions", "userId、projectId", "保存用户的机器级权限"),
        ("user_program_permissions", "userId、programId", "保存用户的程序级权限"),
        ("api_tokens", "userId、name、tokenHash、scopes、expiresAt、revokedAt", "保存服务令牌元数据和摘要"),
        ("api_audit_events", "tokenId、userId、method、path、status、duration、ip", "保存服务API调用审计"),
    ], [2200, 3800, 3360], font_size=8.9)
    add_heading(doc, "6.3 关键约束", 2)
    add_list(doc, [
        "用户名、项目名、分组名和令牌摘要等需要唯一性的字段由数据库约束保证。",
        "用户删除时，其项目权限和程序权限通过外键级联清理。",
        "分组删除时，项目的groupId置空，避免误删除机器记录。",
        "Supervisor密码不以明文保存，读取配置时由服务端解密后使用。",
        "API令牌只保存摘要，创建响应之外不再提供令牌明文。",
    ])
    add_heading(doc, "6.4 程序标识", 2)
    add_body(doc, "系统使用可解析的程序标识关联机器和程序名称。后端programId工具负责在接口路径和权限记录中解析该标识，使同名程序在不同机器上仍可准确定位。程序级权限记录使用完整程序标识，权限校验时先判断该记录，再回退到机器级权限。")

    add_heading(doc, "7 用户操作说明", 1)
    add_heading(doc, "7.1 首次启动与管理员账户", 2)
    add_body(doc, "系统第一次连接空数据库时自动创建角色和超级管理员admin，并生成随机初始密码。管理员应从受控的启动日志中获取该密码，首次登录后立即修改。生产环境不得在工单、聊天记录或公开文档中长期保存初始密码。")
    add_heading(doc, "7.2 登录系统", 2)
    add_list(doc, [
        "在浏览器中打开系统地址，例如生产单端口部署的http(s)://服务器地址:6002。",
        "在登录页输入用户名和密码。",
        "单击“登录”按钮。系统验证成功后进入进程管理主面板。",
        "若连续多次输入错误，系统可能触发每IP每分钟5次的登录限制，应稍后重试并核对账户状态。",
    ], numbered=True)
    add_heading(doc, "7.3 主面板布局", 2)
    add_body(doc, "主面板由左侧资源树、顶部工具区、统计区和进程列表组成。左侧资源树按分组展示机器，并显示连接状态；顶部区域提供搜索、折叠侧栏、用户菜单和管理入口；统计区显示程序总数、运行中、已停止和异常数；列表区域展示当前选择范围内的程序及其操作按钮。")
    add_table(doc, ["区域", "显示内容", "常用操作"], [
        ("资源树", "分组、机器、连接状态", "展开分组、选择机器、搜索机器"),
        ("统计区", "总数、运行、停止、异常", "快速判断运行态势"),
        ("进程表", "程序名、状态、运行时间、操作", "搜索、启动、停止、重启、查看日志"),
        ("用户菜单", "当前用户与管理入口", "修改密码、用户管理、退出登录"),
    ], [1700, 3700, 3960])
    add_heading(doc, "7.4 查看机器和进程", 2)
    add_list(doc, [
        "在左侧分组树中展开目标分组。",
        "观察机器旁的状态标识；连接检查中、在线和离线采用不同状态展示。",
        "单击目标机器，页面加载该机器的进程列表。",
        "在搜索框输入程序名关键字，列表即时过滤。",
        "查看状态标签和运行时间；异常状态包括FATAL、BACKOFF、UNKNOWN和EXITED等。",
    ], numbered=True)
    add_heading(doc, "7.5 控制单个进程", 2)
    add_list(doc, [
        "确认当前选择的机器和程序名称，避免操作错误目标。",
        "根据当前状态单击启动、停止或重启图标按钮。",
        "等待系统显示“指令已发送”或错误提示。",
        "系统在操作后刷新状态；若Supervisor执行需要时间，可稍后观察状态变化。",
        "若收到403提示，说明当前账户缺少该机器或程序权限，应联系管理员调整授权。",
    ], numbered=True)
    add_note(doc, "操作风险", "停止、重启和全部操作会影响业务进程可用性。生产环境应先确认维护窗口、目标机器和程序名称，并遵循公司变更管理制度。", fill="F4EDED", color=RED)
    add_heading(doc, "7.6 批量控制与配置重载", 2)
    add_body(doc, "在机器级操作菜单中，可选择全部启动、全部停止、全部重启和重载Supervisor配置。批量操作作用于当前机器上的全部程序，影响范围大。执行前应检查机器名称、程序列表和业务窗口；重载配置用于让Supervisor重新读取配置，不等同于修改SuperD中的机器信息。")
    add_heading(doc, "7.7 查看实时日志", 2)
    add_list(doc, [
        "在进程列表中进入目标程序的详情或日志入口。",
        "选择stdout或stderr日志类型。",
        "日志终端建立Socket.io连接后显示新的增量输出。",
        "需要检查历史内容时可暂停滚动；恢复后继续接收后续日志。",
        "选择终端文本即可复制，注意避免将密钥、令牌或个人信息粘贴到不受控位置。",
    ], numbered=True)
    add_heading(doc, "7.8 管理分组和机器", 2)
    add_body(doc, "超级管理员通过项目管理弹窗新增、编辑或删除机器，并设置名称、说明、Supervisor主机、端口、用户名、密码和所属分组。新增后系统检查连接状态。编辑密码时应使用受控凭据；删除机器只删除SuperD中的连接记录，不会自动卸载目标服务器的Supervisor或删除目标进程。")
    add_heading(doc, "7.9 管理用户", 2)
    add_body(doc, "超级管理员可查看全部管理员和普通用户，并调整角色、上级管理员、密码和资源权限。普通管理员只能管理自己的组员，并只能创建普通用户。系统界面支持按用户名、角色或上级关系筛选用户。admin账号受特殊保护，不应被普通管理操作删除或降级。")
    add_heading(doc, "7.10 配置资源权限", 2)
    add_list(doc, [
        "进入用户管理并选择目标普通用户。",
        "若用户需要操作某台机器的全部程序，为其添加机器级权限。",
        "若用户只需要操作少量程序，为其添加具体程序权限。",
        "保存后使用该用户账户验证资源可见范围和操作结果。",
        "移除权限后再次验证，确保用户无法通过接口绕过前端访问未授权资源。",
    ], numbered=True)
    add_heading(doc, "7.11 修改密码与退出", 2)
    add_body(doc, "用户可从个人菜单打开修改密码窗口，按系统要求输入新密码并提交。管理员也可在权限范围内为受管用户重置密码。完成操作后应退出登录，尤其是在共享终端或临时运维电脑上。退出后前端清理登录态，访问受保护页面会重新跳转到登录页。")

    add_heading(doc, "8 服务API使用说明", 1)
    add_heading(doc, "8.1 适用场景", 2)
    add_body(doc, "服务API用于告警平台、巡检任务、发布系统或其他长期运行程序访问SuperD。网页登录返回的JWT适合人员交互，不应作为长期集成凭据；外部服务应使用专门创建的API Token，并只授予所需scope和资源权限。")
    add_heading(doc, "8.2 创建与保管令牌", 2)
    add_list(doc, [
        "超级管理员先创建或选择一个专用于集成的用户，并配置必要的机器/程序权限。",
        "使用管理员JWT调用POST /api/api-tokens，提交userId、name、scopes和有效期。",
        "创建响应只返回一次sd_开头的明文令牌，应立即写入调用方的密钥管理系统。",
        "不得把令牌提交到Git仓库、普通日志、聊天群或公开文档。",
        "轮换时先创建新令牌并完成调用方切换，再撤销旧令牌。",
    ], numbered=True)
    add_heading(doc, "8.3 请求方式", 2)
    add_body(doc, "调用方在HTTP Authorization请求头中使用Bearer方案携带服务令牌，例如“Authorization: Bearer sd_<SERVICE_TOKEN>”。接口以/api为稳定v1基线，可通过GET /api/version查询版本。调用方应依据HTTP状态码和message字段处理失败，不依赖未文档化字段。")
    add_heading(doc, "8.4 Scope范围", 2)
    add_table(doc, ["Scope", "允许的能力"], [
        ("projects:read / projects:write", "读取或维护机器记录"),
        ("groups:read / groups:write", "读取或维护项目分组"),
        ("programs:read / programs:write", "读取进程或执行进程控制"),
        ("logs:read", "读取stdout/stderr日志"),
        ("users:read / users:write", "读取或维护用户和权限"),
        ("tokens:manage", "管理服务令牌和查看相关审计"),
        ("*", "全部scope，仅限严格受控的管理员服务"),
    ], [3300, 6060])
    add_heading(doc, "8.5 撤销与审计", 2)
    add_body(doc, "管理员可通过DELETE /api/api-tokens/:id立即撤销令牌。服务令牌访问受保护接口时，系统在响应完成后写入审计事件，可通过GET /api/api-audit-events查询最近记录。审计记录有助于定位调用方、路径、结果、耗时和来源IP，但企业仍应结合反向代理、主机和数据库日志建立完整留痕。")

    add_heading(doc, "9 安全设计", 1)
    add_heading(doc, "9.1 身份与口令安全", 2)
    add_body(doc, "用户密码使用bcrypt哈希保存，系统不需要恢复原始密码。网页认证使用JWT，服务集成使用独立API Token。初始管理员密码随机生成，首次登录后应立即变更。安全密钥不得硬编码到源代码，通过环境变量或受控的.env文件注入。")
    add_heading(doc, "9.2 连接凭据保护", 2)
    add_body(doc, "Supervisor连接信息中的密码在数据库中加密保存，调用XML-RPC前在服务端解密。生产环境应确保ENCRYPTION_KEY具备足够随机性并纳入密钥管理；更换该密钥前必须制定已存凭据迁移方案，否则旧密文可能无法解密。")
    add_heading(doc, "9.3 权限判定", 2)
    add_body(doc, "每个受保护请求先完成身份验证，再根据功能检查角色、API scope和资源权限。普通用户的程序级权限优先，机器级权限兜底；普通管理员的用户管理范围受createdBy关系限制；全局机器和分组配置仅允许超级管理员操作。前端隐藏无权功能仅用于改善体验，最终权限以服务端中间件判定为准。")
    add_heading(doc, "9.4 Web安全", 2)
    add_list(doc, [
        "使用Helmet设置常见HTTP安全头，内容安全策略按前端部署需要单独控制。",
        "Session Cookie启用httpOnly和sameSite=lax，生产模式启用secure。",
        "登录接口设置速率限制，降低暴力尝试风险。",
        "CORS仅允许配置来源或同源请求，并允许受控凭据。",
        "请求体大小限制为10MB，降低异常超大请求影响。",
    ])
    add_heading(doc, "9.5 安全运维建议", 2)
    add_list(doc, [
        "在反向代理层启用HTTPS、访问控制、请求限速和安全日志。",
        "限制Supervisor 9001端口只接受SuperD应用服务器访问。",
        "定期轮换JWT、Session、加密密钥和服务API令牌。",
        "按照最小权限原则配置机器/程序授权和令牌scope。",
        "定期备份数据库并验证恢复，备份文件同样包含敏感配置，应加密保管。",
        "上线前运行依赖漏洞扫描、接口权限测试和配置审计。",
    ])

    add_heading(doc, "10 异常处理与日志", 1)
    add_heading(doc, "10.1 统一错误处理", 2)
    add_body(doc, "后端通过ApiError和统一错误处理中间件生成一致的HTTP错误响应。未匹配的API路由返回404；认证失败通常返回401并触发前端跳转登录；权限不足返回403；参数或业务状态错误返回4xx；未预期异常记录详细错误后返回受控的5xx响应。前端API封装避免重复弹出已统一处理的错误。")
    add_heading(doc, "10.2 常见异常", 2)
    add_table(doc, ["现象", "可能原因", "处理建议"], [
        ("无法登录", "密码错误、速率限制、JWT密钥配置异常", "核对账号；等待限流窗口；检查服务日志和密钥配置"),
        ("机器显示离线", "网络不通、端口未开放、Supervisor未启动、凭据错误", "从应用服务器测试网络和XML-RPC配置"),
        ("进程操作返回403", "用户无机器/程序权限或令牌scope不足", "检查绑定用户权限和令牌scope交集"),
        ("日志无内容", "日志文件不存在、程序未输出、连接未建立或已降频", "核对Supervisor日志配置和Socket.io连接"),
        ("状态短暂未变化", "Supervisor执行需要时间或缓存尚未刷新", "等待下一次刷新；检查目标程序实际状态"),
        ("数据库启动失败", "文件权限、MySQL连接、密码或表结构问题", "检查存储变量、目录权限和数据库服务"),
    ], [2100, 3500, 3760], font_size=8.8)
    add_heading(doc, "10.3 日志分类", 2)
    add_body(doc, "应用Logger记录请求、响应、调试、信息、警告、错误和致命事件，并写入访问日志和错误日志。服务API调用另有结构化审计记录。生产环境应设置日志轮转、保留期限、访问权限和脱敏规则，避免JWT、API Token、密码等敏感信息进入日志。PM2配置可用于进程守护和日志文件管理。")

    add_heading(doc, "11 安装部署与维护", 1)
    add_heading(doc, "11.1 安装准备", 2)
    add_list(doc, [
        "准备Node.js运行环境、npm依赖安装能力和PM2。",
        "准备SQLite可写目录，或预先创建MySQL数据库和专用账户。",
        "在所有目标服务器安装并配置Supervisor inet_http_server。",
        "规划SuperD应用服务器到Supervisor端口的网络访问策略。",
        "生成SESSION_SECRET、JWT_SECRET和ENCRYPTION_KEY并存入受控配置。",
    ], numbered=True)
    add_heading(doc, "11.2 开发模式启动", 2)
    add_body(doc, "开发模式下，后端进入backend目录安装依赖并运行npm run dev，默认监听3000端口；前端进入client目录安装依赖并运行npm run dev，默认监听5173端口。开发模式使用分端口结构，CORS配置允许指定的本地前端来源。首次启动会初始化数据库并创建管理员。")
    add_heading(doc, "11.3 生产模式部署", 2)
    add_list(doc, [
        "在client目录执行npm run build，生成前端dist静态文件。",
        "在backend目录安装生产依赖并配置环境变量。",
        "根据需要选择SQLite或MySQL，并验证应用账户具有必要权限。",
        "使用PM2启动ecosystem.config.js，示例配置在6002端口提供单入口服务。",
        "在反向代理上配置HTTPS，并限制后台和Supervisor端口的访问范围。",
        "使用管理员初始密码登录，立即修改密码并新增受控运维账户。",
    ], numbered=True)
    add_heading(doc, "11.4 Supervisor端配置", 2)
    add_body(doc, "每台被管理服务器需在Supervisor配置中启用inet_http_server，设置监听地址、端口、用户名和密码。出于安全考虑，不应将该端口直接暴露到公网；应通过主机防火墙、安全组或专用网络，仅允许SuperD应用服务器访问。配置变更后重载或重启Supervisor，并在SuperD中新增机器记录进行连接测试。")
    add_heading(doc, "11.5 备份与恢复", 2)
    add_body(doc, "SQLite模式应定期备份数据库文件，并在写入活动较低或使用数据库一致性备份机制时执行；MySQL模式应使用企业现有数据库备份策略。恢复演练要验证用户、权限、机器、分组和令牌元数据的完整性。加密密钥必须与数据库备份关联保管，否则恢复后的Supervisor密码可能无法解密。")
    add_heading(doc, "11.6 升级维护", 2)
    add_body(doc, "升级前应备份数据库与配置，阅读版本说明并在测试环境验证前端构建、数据库兼容、登录、权限、进程控制和日志功能。升级过程中保持软件名称和版本号管理一致；若用于软著新版本登记，应另行确认版本功能变化、完成日期和权利信息。升级后检查/api/version、关键接口、Socket.io连接和PM2运行状态。")

    add_heading(doc, "12 测试与验收建议", 1)
    add_heading(doc, "12.1 功能测试", 2)
    add_table(doc, ["测试域", "关键检查点", "期望结果"], [
        ("登录认证", "正确/错误密码、限流、退出、过期登录态", "身份结果正确，未登录无法访问受保护页面"),
        ("机器与分组", "增删改、分组移动、连接状态、自然排序", "数据持久化且资源树展示一致"),
        ("进程查询", "单机/全局列表、搜索、10秒刷新、状态统计", "列表无重复，状态和统计一致"),
        ("进程控制", "启动、停止、重启、批量操作、重载", "权限正确，操作结果和状态更新一致"),
        ("实时日志", "stdout/stderr、暂停、复制、无日志降频", "增量输出连续且资源消耗受控"),
        ("用户权限", "三级角色、上级关系、机器/程序权限", "越权请求返回403，授权资源可正常操作"),
        ("服务令牌", "创建、scope、过期、撤销、审计", "明文只出现一次，摘要存储，审计完整"),
    ], [1700, 3900, 3760], font_size=8.5)
    add_heading(doc, "12.2 安全测试", 2)
    add_list(doc, [
        "验证未携带、伪造或过期JWT/API Token时的拒绝行为。",
        "验证普通用户不能调用用户管理、全局分组和机器管理接口。",
        "验证令牌scope不足、绑定用户资源权限不足时均返回拒绝。",
        "检查日志和错误响应中是否泄露密码、JWT、令牌明文或数据库摘要。",
        "检查CORS、Cookie、HTTPS反向代理和Supervisor端口访问策略。",
        "执行依赖漏洞、输入校验、速率限制和常见Web安全测试。",
    ])
    add_heading(doc, "12.3 兼容性与性能测试", 2)
    add_body(doc, "应在企业计划支持的浏览器、Node.js版本和数据库模式上执行兼容性测试。性能测试应覆盖多机器并发状态查询、10秒自动刷新、批量控制、长时间日志连接和审计写入。重点观察XML-RPC响应、5秒缓存命中率、Socket.io连接数、内存使用和数据库延迟，并设置符合企业规模的容量指标。")
    add_heading(doc, "12.4 验收标准建议", 2)
    add_list(doc, [
        "安装和初始化流程可重复执行，必要环境变量有明确校验。",
        "用户可按角色和资源权限完成预期操作，不能通过直接接口越权。",
        "机器状态、进程列表、统计和操作后状态在约定时间内保持一致。",
        "日志可持续接收且断开后能重新连接，无新日志时能够自动降频。",
        "数据库备份可恢复，密钥和凭据管理符合公司安全要求。",
        "使用说明、接口文档、版本号和最终交存代码保持一致。",
    ])

    add_heading(doc, "13 技术特点与独创性说明", 1)
    add_body(doc, "SuperD围绕多服务器Supervisor运维场景形成了统一的Web管理模型，将分散的命令行控制转换为带身份、资源层级、实时状态和可视化反馈的浏览器工作流。软件的技术特点不是单一调用Supervisor接口，而是对多机资源组织、权限粒度、缓存一致性、实时日志和服务集成安全进行组合设计。")
    add_heading(doc, "13.1 多层级资源组织", 2)
    add_body(doc, "系统把分组、机器和程序构成统一资源树，并对机器名称进行前缀与尾号自然排序。用户可以从分组定位机器，再进入具体程序，同时支持全局汇总和关键字搜索。该模型适应机器数量增长和多业务团队协作。")
    add_heading(doc, "13.2 双粒度权限模型", 2)
    add_body(doc, "在三级角色基础上，系统进一步提供机器级和程序级授权。程序级权限先于机器级权限判定，既支持整机运维，也支持只授权单个程序；普通管理员通过上级关系管理自己的组员，形成角色边界与资源边界的组合控制。")
    add_heading(doc, "13.3 状态与日志的自适应机制", 2)
    add_body(doc, "进程查询采用短周期缓存，界面采用10秒静默刷新，控制操作后主动清除缓存。实时日志按偏移量读取，只推送新增内容，并根据空日志次数在1秒和10秒轮询之间自适应切换。该组合在及时性、界面稳定性和服务器负载之间取得平衡。")
    add_heading(doc, "13.4 人员访问与服务访问分离", 2)
    add_body(doc, "网页登录使用JWT，长期服务调用使用独立API Token。服务令牌具有scope、有效期、撤销状态、摘要存储和调用审计，并且还要受到绑定用户资源权限限制。两套认证方式对应不同使用主体，避免把人员短期令牌直接用于自动化服务。")

    add_heading(doc, "附录A 主要接口清单", 1)
    api_rows = [
        ("认证", "POST", "/api/login", "用户登录"), ("认证", "GET", "/api/user", "获取当前用户"),
        ("机器", "GET", "/api/projects", "获取可见机器"), ("机器", "POST", "/api/projects", "创建机器"),
        ("机器", "PUT", "/api/projects/:id", "更新机器"), ("机器", "DELETE", "/api/projects/:id", "删除机器"),
        ("机器", "GET", "/api/projects/:id/status", "检查连接状态"),
        ("分组", "GET", "/api/groups", "分组列表"), ("分组", "POST", "/api/groups", "创建分组"),
        ("分组", "PUT", "/api/groups/:id", "更新分组"), ("分组", "DELETE", "/api/groups/:id", "删除分组"),
        ("进程", "GET", "/api/projects/:id/programs", "机器进程列表"), ("进程", "GET", "/api/programs", "全部可见进程"),
        ("进程", "POST", "/api/programs/:id/start", "启动进程"), ("进程", "POST", "/api/programs/:id/stop", "停止进程"),
        ("进程", "POST", "/api/programs/:id/restart", "重启进程"), ("进程", "GET", "/api/programs/:id/stdout", "读取标准输出"),
        ("进程", "GET", "/api/programs/:id/stderr", "读取错误输出"),
        ("批量", "POST", "/api/projects/:id/programs/start-all", "全部启动"),
        ("批量", "POST", "/api/projects/:id/programs/stop-all", "全部停止"),
        ("批量", "POST", "/api/projects/:id/programs/restart-all", "全部重启"),
        ("批量", "POST", "/api/projects/:id/reload", "重载Supervisor配置"),
        ("用户", "GET", "/api/users", "用户列表"), ("用户", "POST", "/api/users", "创建用户"),
        ("用户", "PUT", "/api/users/:id/role", "修改角色"), ("用户", "PUT", "/api/users/:id/password", "管理员重置密码"),
        ("权限", "GET/POST/DELETE", "/api/users/:id/project-permissions", "机器权限管理"),
        ("权限", "GET/POST/DELETE", "/api/users/:id/program-permissions", "程序权限管理"),
        ("令牌", "GET/POST", "/api/api-tokens", "查询/创建服务令牌"),
        ("令牌", "DELETE", "/api/api-tokens/:id", "撤销服务令牌"),
        ("审计", "GET", "/api/api-audit-events", "查询服务API审计"),
        ("版本", "GET", "/api/version", "查询API版本"),
    ]
    add_table(doc, ["模块", "方法", "路径", "说明"], api_rows, [1200, 1500, 4500, 2160], font_size=7.8,
              aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])

    add_heading(doc, "附录B 功能与源代码对应关系", 1)
    mapping_rows = [
        ("应用启动与中间件", "backend/app.js", "环境校验、静态托管、路由、Session、管理员初始化、Socket.io启动"),
        ("全局配置", "backend/config.js", "端口、存储、Supervisor、日志和安全配置"),
        ("数据库初始化", "backend/init-db.js", "SQLite/MySQL表结构、角色初始化"),
        ("认证与权限", "backend/middleware/auth.js", "JWT/API Token认证、角色、机器与程序权限"),
        ("用户认证接口", "backend/routes/auth.js", "登录、当前用户、退出相关处理"),
        ("机器管理", "backend/routes/projects.js", "机器CRUD、连接状态、分组设置"),
        ("分组管理", "backend/routes/groups.js", "分组CRUD及分组机器查询"),
        ("进程与日志接口", "backend/routes/programs.js", "列表、详情、启停重启、批量、日志和重载"),
        ("用户与权限", "backend/routes/users.js", "用户、角色、密码、上级关系、资源权限"),
        ("服务令牌", "backend/routes/apiTokens.js", "令牌创建、查询、撤销和审计查询"),
        ("Supervisor适配", "backend/services/supervisorService.js", "XML-RPC调用、进程缓存、日志读取"),
        ("实时日志", "backend/services/socketServer.js", "Socket.io连接、日志偏移、推送和自适应轮询"),
        ("SQLite实现", "backend/models/db.sqlite.js", "SQLite数据访问方法"),
        ("MySQL实现", "backend/models/db.mysql.js", "MySQL数据访问方法"),
        ("加密与令牌", "backend/utils/crypto.js / apiToken.js", "Supervisor密码加解密、服务令牌生成与摘要"),
        ("前端路由", "client/src/App.jsx", "认证上下文、受保护路由和页面入口"),
        ("登录页面", "client/src/pages/LoginPage.jsx", "登录表单和错误反馈"),
        ("进程主面板", "client/src/pages/ProgramsPage.jsx", "资源树、统计、列表、搜索、控制和自动刷新"),
        ("进程详情", "client/src/pages/ProgramDetailPage.jsx", "进程详情与日志抽屉"),
        ("日志终端", "client/src/components/LogTerminal.jsx", "xterm.js显示、暂停和Socket.io日志"),
        ("用户管理", "client/src/pages/UsersPage.jsx", "用户分组、筛选和管理"),
        ("权限编辑", "client/src/components/users/*", "用户表格、表单和资源权限配置"),
        ("API封装", "client/src/utils/api.js", "统一HTTP请求和接口函数"),
    ]
    add_table(doc, ["功能", "源文件", "对应说明"], mapping_rows, [2100, 3600, 3660], font_size=7.7)

    add_heading(doc, "附录C 界面视图清单", 1)
    add_body(doc, "系统各界面的构成与主要交互元素如下：")
    add_table(doc, ["界面", "主要构成", "核心交互"], [
        ("登录页", "软件标识、用户名/密码表单", "登录校验、失败提示、跳转主面板"),
        ("进程管理主面板", "分组机器树、状态统计、进程列表与搜索", "按分组/机器浏览进程、搜索、批量启停重启"),
        ("单机进程列表", "程序状态、运行时间、控制按钮", "启动/停止/重启单个进程、重载Supervisor配置"),
        ("实时日志窗口", "stdout/stderr终端输出、暂停与复制控件", "按偏移量增量推送日志、暂停滚动、复制内容"),
        ("机器管理弹窗", "机器信息、Supervisor连接配置、分组设置", "机器增删改、连接状态检查、分组调整"),
        ("用户管理页面", "三级角色、上级管理员、用户列表", "用户创建、角色分配、密码重置、禁用"),
        ("资源权限配置", "机器级与程序级权限设置", "按机器或单个程序授权、权限边界组合"),
        ("服务令牌管理", "令牌名称、scope、到期与撤销状态", "令牌创建/撤销，明文仅创建时展示一次"),
        ("API审计记录", "请求方法、路径、状态码、耗时与来源", "服务API调用流水查询"),
    ], [1900, 3300, 4160], font_size=8.8)

    add_heading(doc, "结束语", 1)
    add_body(doc, "SuperD通过Web化资源管理、三级角色与双粒度资源授权、XML-RPC进程控制、实时日志推送、短周期缓存和服务API安全机制，为多服务器Supervisor运维提供统一、可控和可扩展的管理平台。本说明书覆盖软件的技术设计和用户操作，可作为软件著作权登记中的文档鉴别材料。")

    # Document core properties contain no personal metadata.
    props = doc.core_properties
    props.title = f"{SOFTWARE_NAME} {VERSION} 软件著作权登记说明书"
    props.subject = "计算机软件著作权登记文档鉴别材料"
    props.author = RIGHTSHOLDER
    props.keywords = "SuperD, Supervisor, 进程管理, 软件著作权"
    props.comments = "基于项目仓库生成；提交前需核定申请人主体信息。"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
